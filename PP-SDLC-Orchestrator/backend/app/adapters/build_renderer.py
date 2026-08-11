"""Renders the Build Agent's two artefacts — shared by BuildMockAdapter
and BuildLlmAdapter so the two runtimes differ only in where finding
content comes from, not in how the templates get filled.
"""

import hashlib
from pathlib import Path

from docx import Document

from app.agents_registry.contract import ProducedArtefact

BUILD_REVIEW_ARTEFACT_TYPE = "build_review_report"
FINAL_CODE_REVIEW_ARTEFACT_TYPE = "final_code_review_report"
BUILD_REVIEW_TEMPLATE_RELATIVE_PATH = "04_Templates/build_review_report.docx"
FINAL_CODE_REVIEW_TEMPLATE_RELATIVE_PATH = "04_Templates/final_code_review_report.docx"


def render_build_review(
    *,
    repo_root: Path,
    output_dir: Path,
    project_name: str,
    version_label: str,
    implementation_assets_text: str,
    configuration_summary_text: str,
    findings: list[str],
    defect_entities: list[str],
    fixes_applied: list[str],
) -> ProducedArtefact:
    doc = Document(str(repo_root / BUILD_REVIEW_TEMPLATE_RELATIVE_PATH))
    finding_lines = [f"{eid}: {text}" for eid, text in zip(defect_entities, findings, strict=True)]
    for para in doc.paragraphs:
        if "{{PROJECT_NAME}}" in para.text:
            para.text = para.text.replace("{{PROJECT_NAME}}", str(project_name))
        elif "{{VERSION_LABEL}}" in para.text:
            para.text = para.text.replace("{{VERSION_LABEL}}", version_label)
        elif "{{IMPLEMENTATION_ASSETS}}" in para.text:
            para.text = implementation_assets_text
        elif "{{CONFIGURATION_SUMMARY}}" in para.text:
            para.text = configuration_summary_text
        elif "{{BUILD_FINDINGS}}" in para.text:
            para.text = ""
            for line in finding_lines:
                doc.add_paragraph(line)
        elif "{{FIXES_APPLIED}}" in para.text:
            para.text = ""
            for line in fixes_applied:
                doc.add_paragraph(line)

    output_path_dir = output_dir / BUILD_REVIEW_ARTEFACT_TYPE
    output_path_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_path_dir / f"{version_label}.docx"
    doc.save(output_path)
    checksum = hashlib.sha256(output_path.read_bytes()).hexdigest()

    return ProducedArtefact(
        artefact_type=BUILD_REVIEW_ARTEFACT_TYPE,
        stable_key=BUILD_REVIEW_ARTEFACT_TYPE,
        file_path=str(output_path),
        checksum=checksum,
        entities=list(defect_entities),
    )


def render_final_code_review(
    *,
    repo_root: Path,
    output_dir: Path,
    project_name: str,
    version_label: str,
    review_scope_text: str,
    findings_text: str,
    defect_entities: list[str],
    resolution_lines: list[str],
) -> ProducedArtefact:
    doc = Document(str(repo_root / FINAL_CODE_REVIEW_TEMPLATE_RELATIVE_PATH))
    for para in doc.paragraphs:
        if "{{PROJECT_NAME}}" in para.text:
            para.text = para.text.replace("{{PROJECT_NAME}}", str(project_name))
        elif "{{VERSION_LABEL}}" in para.text:
            para.text = para.text.replace("{{VERSION_LABEL}}", version_label)
        elif "{{REVIEW_SCOPE}}" in para.text:
            para.text = review_scope_text
        elif "{{FINDINGS}}" in para.text:
            para.text = findings_text
        elif "{{RESOLUTION_STATUS}}" in para.text:
            para.text = ""
            for line in resolution_lines:
                doc.add_paragraph(line)

    output_path_dir = output_dir / FINAL_CODE_REVIEW_ARTEFACT_TYPE
    output_path_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_path_dir / f"{version_label}.docx"
    doc.save(output_path)
    checksum = hashlib.sha256(output_path.read_bytes()).hexdigest()

    return ProducedArtefact(
        artefact_type=FINAL_CODE_REVIEW_ARTEFACT_TYPE,
        stable_key=FINAL_CODE_REVIEW_ARTEFACT_TYPE,
        file_path=str(output_path),
        checksum=checksum,
        entities=list(defect_entities),
    )
