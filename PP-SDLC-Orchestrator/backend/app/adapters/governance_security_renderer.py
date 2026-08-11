"""Renders the Governance & Security Agent's Governance Document — shared
by GovernanceSecurityMockAdapter and GovernanceSecurityLlmAdapter so the
two runtimes differ only in where the governance content comes from, not
in how the template gets filled.
"""

import hashlib
from pathlib import Path

from docx import Document

from app.agents_registry.contract import ProducedArtefact

ARTEFACT_TYPE = "governance_document"
TEMPLATE_RELATIVE_PATH = "04_Templates/governance_document.docx"


def render(
    *,
    repo_root: Path,
    output_dir: Path,
    project_name: str,
    version_label: str,
    identity_design_text: str,
    permissions_text: str,
    environment_strategy_text: str,
    dlp_lines: list[str],
    connector_governance_text: str,
    licensing_lines: list[str],
    compliance_text: str,
    operational_ownership_text: str,
    audit_requirements_text: str,
) -> ProducedArtefact:
    doc = Document(str(repo_root / TEMPLATE_RELATIVE_PATH))
    for para in doc.paragraphs:
        if "{{PROJECT_NAME}}" in para.text:
            para.text = para.text.replace("{{PROJECT_NAME}}", str(project_name))
        elif "{{VERSION_LABEL}}" in para.text:
            para.text = para.text.replace("{{VERSION_LABEL}}", version_label)
        elif "{{IDENTITY_DESIGN}}" in para.text:
            para.text = identity_design_text
        elif "{{PERMISSIONS}}" in para.text:
            para.text = permissions_text
        elif "{{ENVIRONMENT_STRATEGY}}" in para.text:
            para.text = environment_strategy_text
        elif "{{DLP}}" in para.text:
            para.text = ""
            for line in dlp_lines:
                doc.add_paragraph(line)
        elif "{{CONNECTOR_GOVERNANCE}}" in para.text:
            para.text = connector_governance_text
        elif "{{LICENSING}}" in para.text:
            para.text = ""
            for line in licensing_lines:
                doc.add_paragraph(line)
        elif "{{COMPLIANCE}}" in para.text:
            para.text = compliance_text
        elif "{{OPERATIONAL_OWNERSHIP}}" in para.text:
            para.text = operational_ownership_text
        elif "{{AUDIT_REQUIREMENTS}}" in para.text:
            para.text = audit_requirements_text

    output_path_dir = output_dir / ARTEFACT_TYPE
    output_path_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_path_dir / f"{version_label}.docx"
    doc.save(output_path)
    checksum = hashlib.sha256(output_path.read_bytes()).hexdigest()

    return ProducedArtefact(
        artefact_type=ARTEFACT_TYPE,
        stable_key=ARTEFACT_TYPE,
        file_path=str(output_path),
        checksum=checksum,
        entities=[],
    )
