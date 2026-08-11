"""Renders the Technical Design Agent's two artefacts — shared by
TechnicalDesignMockAdapter and TechnicalDesignLlmAdapter so the two
runtimes differ only in where options/decisions/risks content comes from,
not in how the templates get filled.
"""

import hashlib
from pathlib import Path

from docx import Document

from app.agents_registry.contract import ProducedArtefact

SOLUTION_APPROACH_ARTEFACT_TYPE = "solution_approach"
ARCHITECTURE_HANDBOOK_ARTEFACT_TYPE = "architecture_handbook"
SOLUTION_APPROACH_TEMPLATE_RELATIVE_PATH = "04_Templates/solution_approach.docx"
ARCHITECTURE_HANDBOOK_TEMPLATE_RELATIVE_PATH = "04_Templates/architecture_handbook.docx"


def render_solution_approach(
    *,
    repo_root: Path,
    output_dir: Path,
    project_name: str,
    version_label: str,
    options: list[tuple[str, str]],
    decisions: list[str],
    decision_entities: list[str],
    risks: list[str],
    limitations: list[str],
    dependencies: list[str],
) -> ProducedArtefact:
    doc = Document(str(repo_root / SOLUTION_APPROACH_TEMPLATE_RELATIVE_PATH))
    decision_lines = [f"{eid}: {text}" for eid, text in zip(decision_entities, decisions, strict=True)]
    for para in doc.paragraphs:
        if "{{PROJECT_NAME}}" in para.text:
            para.text = para.text.replace("{{PROJECT_NAME}}", str(project_name))
        elif "{{VERSION_LABEL}}" in para.text:
            para.text = para.text.replace("{{VERSION_LABEL}}", version_label)
        elif "{{OPTION_ANALYSIS}}" in para.text:
            para.text = ""
            for name, tradeoff in options:
                doc.add_paragraph(f"Option — {name}: {tradeoff}")
            doc.add_paragraph(f"Recommended: {options[0][0]}")
        elif "{{ARCHITECTURE_DECISIONS}}" in para.text:
            para.text = ""
            for line in decision_lines:
                doc.add_paragraph(line)
        elif "{{RISKS}}" in para.text:
            para.text = ""
            for risk in risks:
                doc.add_paragraph(risk)
        elif "{{LIMITATIONS}}" in para.text:
            para.text = ""
            for limitation in limitations:
                doc.add_paragraph(limitation)
        elif "{{DEPENDENCIES}}" in para.text:
            para.text = ""
            for dependency in dependencies:
                doc.add_paragraph(dependency)

    output_path_dir = output_dir / SOLUTION_APPROACH_ARTEFACT_TYPE
    output_path_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_path_dir / f"{version_label}.docx"
    doc.save(output_path)
    checksum = hashlib.sha256(output_path.read_bytes()).hexdigest()

    return ProducedArtefact(
        artefact_type=SOLUTION_APPROACH_ARTEFACT_TYPE,
        stable_key=SOLUTION_APPROACH_ARTEFACT_TYPE,
        file_path=str(output_path),
        checksum=checksum,
        entities=list(decision_entities),
    )


def render_architecture_handbook(
    *,
    repo_root: Path,
    output_dir: Path,
    project_name: str,
    version_label: str,
    logical_architecture_text: str,
    integration_overview_text: str,
    infrastructure_overview_text: str,
) -> ProducedArtefact:
    doc = Document(str(repo_root / ARCHITECTURE_HANDBOOK_TEMPLATE_RELATIVE_PATH))
    for para in doc.paragraphs:
        if "{{PROJECT_NAME}}" in para.text:
            para.text = para.text.replace("{{PROJECT_NAME}}", str(project_name))
        elif "{{VERSION_LABEL}}" in para.text:
            para.text = para.text.replace("{{VERSION_LABEL}}", version_label)
        elif "{{LOGICAL_ARCHITECTURE}}" in para.text:
            para.text = logical_architecture_text
        elif "{{INTEGRATION_OVERVIEW}}" in para.text:
            para.text = integration_overview_text
        elif "{{INFRASTRUCTURE_OVERVIEW}}" in para.text:
            para.text = infrastructure_overview_text

    output_path_dir = output_dir / ARCHITECTURE_HANDBOOK_ARTEFACT_TYPE
    output_path_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_path_dir / f"{version_label}.docx"
    doc.save(output_path)
    checksum = hashlib.sha256(output_path.read_bytes()).hexdigest()

    return ProducedArtefact(
        artefact_type=ARCHITECTURE_HANDBOOK_ARTEFACT_TYPE,
        stable_key=ARCHITECTURE_HANDBOOK_ARTEFACT_TYPE,
        file_path=str(output_path),
        checksum=checksum,
        entities=[],
    )
