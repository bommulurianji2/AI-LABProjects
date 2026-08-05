"""Renders 04_Templates/requirement_specification.docx — shared by
AnalysisMockAdapter and AnalysisLlmAdapter so the two runtimes differ only
in where scope/requirements/assumptions text comes from, not in how the
template gets filled.
"""

from pathlib import Path

from docx import Document

TEMPLATE_RELATIVE_PATH = "04_Templates/requirement_specification.docx"


def render(
    *,
    repo_root: Path,
    output_path: Path,
    project_name: str,
    version_label: str,
    scope_text: str,
    requirement_lines: list[str],
    assumptions_text: str,
) -> None:
    doc = Document(str(repo_root / TEMPLATE_RELATIVE_PATH))
    for para in doc.paragraphs:
        if "{{PROJECT_NAME}}" in para.text:
            para.text = para.text.replace("{{PROJECT_NAME}}", str(project_name))
        elif "{{VERSION_LABEL}}" in para.text:
            para.text = para.text.replace("{{VERSION_LABEL}}", version_label)
        elif "{{SCOPE}}" in para.text:
            para.text = scope_text
        elif "{{REQUIREMENTS_TABLE}}" in para.text:
            para.text = ""
            for line in requirement_lines:
                doc.add_paragraph(line)
        elif "{{ASSUMPTIONS}}" in para.text:
            para.text = assumptions_text

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
