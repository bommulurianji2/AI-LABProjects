"""Renders the Data & Integration Agent's Data Design Document — shared by
DataIntegrationMockAdapter and DataIntegrationLlmAdapter so the two
runtimes differ only in where schema/relationship content comes from, not
in how the template gets filled.
"""

import hashlib
from pathlib import Path

from docx import Document

from app.agents_registry.contract import ProducedArtefact

ARTEFACT_TYPE = "data_design_document"
TEMPLATE_RELATIVE_PATH = "04_Templates/data_design_document.docx"


def render(
    *,
    repo_root: Path,
    output_dir: Path,
    project_name: str,
    version_label: str,
    entities: list[tuple[str, str]],
    entity_ids: list[str],
    relationships: list[str],
    external_sources: list[str],
    connectors: list[str],
    data_migration_text: str,
    reporting_model_text: str,
) -> ProducedArtefact:
    doc = Document(str(repo_root / TEMPLATE_RELATIVE_PATH))
    entity_lines = [f"{eid}: {name} — {desc}" for eid, (name, desc) in zip(entity_ids, entities, strict=True)]
    for para in doc.paragraphs:
        if "{{PROJECT_NAME}}" in para.text:
            para.text = para.text.replace("{{PROJECT_NAME}}", str(project_name))
        elif "{{VERSION_LABEL}}" in para.text:
            para.text = para.text.replace("{{VERSION_LABEL}}", version_label)
        elif "{{DATAVERSE_SCHEMA}}" in para.text:
            para.text = ""
            for line in entity_lines:
                doc.add_paragraph(line)
        elif "{{RELATIONSHIPS}}" in para.text:
            para.text = ""
            for rel in relationships:
                doc.add_paragraph(rel)
        elif "{{EXTERNAL_SOURCES}}" in para.text:
            para.text = ""
            for src in external_sources:
                doc.add_paragraph(src)
        elif "{{CONNECTORS}}" in para.text:
            para.text = ""
            for conn in connectors:
                doc.add_paragraph(conn)
        elif "{{DATA_MIGRATION}}" in para.text:
            para.text = data_migration_text
        elif "{{REPORTING_MODEL}}" in para.text:
            para.text = reporting_model_text

    output_path_dir = output_dir / ARTEFACT_TYPE
    output_path_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_path_dir / f"{version_label}.docx"
    doc.save(output_path)
    checksum = hashlib.sha256(output_path.read_bytes()).hexdigest()

    return ProducedArtefact(
        artefact_type=ARTEFACT_TYPE,
        stable_key=ARTEFACT_TYPE,
        file_path=str(output_path),
        checksum=checksum,
        entities=list(entity_ids),
    )
