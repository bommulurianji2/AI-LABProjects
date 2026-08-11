"""Renders the UX Design Agent's two artefacts — shared by
UxDesignMockAdapter and UxDesignLlmAdapter so the two runtimes differ only
in where personas/journeys/screens content comes from, not in how the
templates get filled.
"""

import hashlib
import html
from pathlib import Path

from docx import Document

from app.agents_registry.contract import ProducedArtefact

SPEC_ARTEFACT_TYPE = "ux_design_specification"
PROTOTYPE_ARTEFACT_TYPE = "ux_interactive_prototype"
SPEC_TEMPLATE_RELATIVE_PATH = "04_Templates/ux_design_specification.docx"
PROTOTYPE_TEMPLATE_RELATIVE_PATH = "04_Templates/ux_interactive_prototype.html"


def render_spec(
    *,
    repo_root: Path,
    output_dir: Path,
    project_name: str,
    version_label: str,
    personas: list[str],
    journeys: list[str],
    screens: list[tuple[str, str]],
    screen_entities: list[str],
    responsive_behavior_text: str,
    accessibility_text: str,
) -> ProducedArtefact:
    doc = Document(str(repo_root / SPEC_TEMPLATE_RELATIVE_PATH))
    screen_lines = [f"{eid}: {name} — {desc}" for eid, (name, desc) in zip(screen_entities, screens, strict=True)]
    for para in doc.paragraphs:
        if "{{PROJECT_NAME}}" in para.text:
            para.text = para.text.replace("{{PROJECT_NAME}}", str(project_name))
        elif "{{VERSION_LABEL}}" in para.text:
            para.text = para.text.replace("{{VERSION_LABEL}}", version_label)
        elif "{{PERSONAS}}" in para.text:
            para.text = ""
            for persona in personas:
                doc.add_paragraph(persona)
        elif "{{JOURNEYS}}" in para.text:
            para.text = ""
            for journey in journeys:
                doc.add_paragraph(journey)
        elif "{{SCREEN_INVENTORY}}" in para.text:
            para.text = ""
            for line in screen_lines:
                doc.add_paragraph(line)
        elif "{{NAVIGATION}}" in para.text:
            para.text = "Top-level navigation: " + " | ".join(name for name, _ in screens)
        elif "{{RESPONSIVE_BEHAVIOR}}" in para.text:
            para.text = responsive_behavior_text
        elif "{{ACCESSIBILITY}}" in para.text:
            para.text = accessibility_text
        elif "{{PROTOTYPE_REF}}" in para.text:
            para.text = para.text.replace("{{PROTOTYPE_REF}}", PROTOTYPE_ARTEFACT_TYPE)

    spec_dir = output_dir / SPEC_ARTEFACT_TYPE
    spec_dir.mkdir(parents=True, exist_ok=True)
    spec_path = spec_dir / f"{version_label}.docx"
    doc.save(spec_path)
    checksum = hashlib.sha256(spec_path.read_bytes()).hexdigest()

    return ProducedArtefact(
        artefact_type=SPEC_ARTEFACT_TYPE,
        stable_key=SPEC_ARTEFACT_TYPE,
        file_path=str(spec_path),
        checksum=checksum,
        entities=list(screen_entities),
    )


def render_prototype(
    *,
    repo_root: Path,
    output_dir: Path,
    project_name: str,
    version_label: str,
    screens: list[tuple[str, str]],
    screen_entities: list[str],
) -> ProducedArtefact:
    template_text = (repo_root / PROTOTYPE_TEMPLATE_RELATIVE_PATH).read_text(encoding="utf-8")

    # project_name is user-supplied (Project.name via the API) and screen
    # name/description may now come from a real model (previously a fixed
    # pool) — all embedded into an HTML file that may be opened in a
    # browser, so everything gets escaped to avoid the prototype becoming
    # an XSS vector.
    safe_project_name = html.escape(str(project_name))

    nav_links = " ".join(
        f'<a href="#{eid}">{html.escape(name)}</a>' for eid, (name, _) in zip(screen_entities, screens, strict=True)
    )
    screens_html = "\n".join(
        f'<section class="screen" id="{eid}"><h2>{eid}: {html.escape(name)}</h2><p>{html.escape(desc)}</p></section>'
        for eid, (name, desc) in zip(screen_entities, screens, strict=True)
    )
    rendered = (
        template_text.replace("{{PROJECT_NAME}}", safe_project_name)
        .replace("{{VERSION_LABEL}}", version_label)
        .replace("{{NAVIGATION_LINKS}}", nav_links)
        .replace("{{SCREENS}}", screens_html)
    )

    prototype_dir = output_dir / PROTOTYPE_ARTEFACT_TYPE
    prototype_dir.mkdir(parents=True, exist_ok=True)
    prototype_path = prototype_dir / f"{version_label}.html"
    prototype_path.write_text(rendered, encoding="utf-8")
    checksum = hashlib.sha256(prototype_path.read_bytes()).hexdigest()

    return ProducedArtefact(
        artefact_type=PROTOTYPE_ARTEFACT_TYPE,
        stable_key=PROTOTYPE_ARTEFACT_TYPE,
        file_path=str(prototype_path),
        checksum=checksum,
        entities=list(screen_entities),
    )
