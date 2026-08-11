"""Renders the Validation / QA Agent's Validation Report — shared by
ValidationQaMockAdapter and ValidationQaLlmAdapter so the two runtimes
differ only in where finding content comes from, not in how the template
gets filled.
"""

import hashlib
from pathlib import Path

from docx import Document

from app.agents_registry.contract import ProducedArtefact

ARTEFACT_TYPE = "validation_report"
TEMPLATE_RELATIVE_PATH = "04_Templates/validation_report.docx"


def render(
    *,
    repo_root: Path,
    output_dir: Path,
    project_name: str,
    version_label: str,
    validation_scope_text: str,
    standards_assessment_text: str,
    findings: list[str],
    overall_verdict_text: str,
) -> ProducedArtefact:
    doc = Document(str(repo_root / TEMPLATE_RELATIVE_PATH))
    for para in doc.paragraphs:
        if "{{PROJECT_NAME}}" in para.text:
            para.text = para.text.replace("{{PROJECT_NAME}}", str(project_name))
        elif "{{VERSION_LABEL}}" in para.text:
            para.text = para.text.replace("{{VERSION_LABEL}}", version_label)
        elif "{{VALIDATION_SCOPE}}" in para.text:
            para.text = validation_scope_text
        elif "{{STANDARDS_ASSESSMENT}}" in para.text:
            para.text = standards_assessment_text
        elif "{{VALIDATION_FINDINGS}}" in para.text:
            para.text = ""
            for finding in findings:
                doc.add_paragraph(finding)
        elif "{{OVERALL_VERDICT}}" in para.text:
            para.text = overall_verdict_text

    output_path_dir = output_dir / ARTEFACT_TYPE
    output_path_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_path_dir / f"{version_label}.docx"
    doc.save(output_path)
    checksum = hashlib.sha256(output_path.read_bytes()).hexdigest()

    return ProducedArtefact(
        artefact_type=ARTEFACT_TYPE,
        stable_key=ARTEFACT_TYPE,
        file_path=str(output_path),
        checksum=checksum,
        entities=[],
    )
