import json

import pytest
from docx import Document

from app.adapters.llm_agent_adapter import DataIntegrationLlmAdapter
from app.adapters.model_providers.base import ModelProviderError
from app.agents_registry.contract import AgentRunRequest
from tests.fixtures.fake_model_provider import FakeModelProvider

_VALID_RESPONSE = {
    "entities": [
        {"name": "Facility", "description": "A bookable room or desk with capacity and amenities."},
        {"name": "Booking", "description": "A reservation of a Facility by an Employee for a time slot."},
    ],
    "relationships": ["Booking N:1 Facility", "Booking N:1 Employee"],
    "external_sources": ["Azure Active Directory maps to the Employee table via object ID."],
    "connectors": ["Office 365 Outlook connector for calendar invitations."],
    "data_migration": "No legacy data migration is required for the initial release.",
    "reporting_model": "Power BI reporting is deferred until reporting requirements are confirmed.",
}


def _make_request(**overrides) -> AgentRunRequest:
    defaults = dict(
        execution_mode="orchestrated",
        project_id="proj-1",
        invocation_id="inv-1",
        agent_id="data_integration",
        task_id="task-1",
        task_request="Design the data model for the booking portal",
        lifecycle_phase="data_integration",
        constraints={"project_name": "Facilities Booking Portal"},
        run_number=1,
        upstream_artefacts_text={"solution_approach": "ADR-001: Use Dataverse as the system of record."},
    )
    defaults.update(overrides)
    return AgentRunRequest(**defaults)


def test_execute_renders_artefact_from_real_model_content():
    adapter = DataIntegrationLlmAdapter(provider=FakeModelProvider(json.dumps(_VALID_RESPONSE)))

    result = adapter.execute(_make_request())

    assert len(result.artefacts_produced) == 1
    produced = result.artefacts_produced[0]
    assert produced.artefact_type == "data_design_document"
    assert produced.entities == ["DATA-001", "DATA-002"]

    doc = Document(produced.file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Facilities Booking Portal" in full_text
    assert "DATA-001: Facility" in full_text
    assert "Booking N:1 Facility" in full_text


def test_upstream_solution_approach_text_is_included_in_prompt():
    provider = FakeModelProvider(json.dumps(_VALID_RESPONSE))
    adapter = DataIntegrationLlmAdapter(provider=provider)

    adapter.execute(_make_request())

    call = provider.calls[0]
    assert "Data & Integration Agent" in call["system"]
    assert "ADR-001: Use Dataverse as the system of record." in call["user"]


def test_missing_upstream_text_omits_context_section_gracefully():
    provider = FakeModelProvider(json.dumps(_VALID_RESPONSE))
    adapter = DataIntegrationLlmAdapter(provider=provider)

    adapter.execute(_make_request(upstream_artefacts_text={}))

    assert "Approved Solution Approach Document" not in provider.calls[0]["user"]


def test_strips_model_supplied_entity_id_prefix():
    response = dict(_VALID_RESPONSE)
    response["entities"] = [{"name": "DATA-1: Facility", "description": "A bookable room or desk."}]
    adapter = DataIntegrationLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    result = adapter.execute(_make_request())

    doc = Document(result.artefacts_produced[0].file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "DATA-001: Facility" in full_text
    assert "DATA-001: DATA" not in full_text


def test_missing_entities_raises_model_provider_error():
    response = dict(_VALID_RESPONSE)
    response["entities"] = []
    adapter = DataIntegrationLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())


def test_malformed_json_raises_model_provider_error():
    adapter = DataIntegrationLlmAdapter(provider=FakeModelProvider("not json"))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())
