import json

import pytest
from docx import Document

from app.adapters.llm_agent_adapter import HypercareClosureLlmAdapter
from app.adapters.model_providers.base import ModelProviderError
from app.agents_registry.contract import AgentRunRequest
from tests.fixtures.fake_model_provider import FakeModelProvider

_VALID_RESPONSE = {
    "hypercare_plan": "Two-week hypercare window post-deployment with daily issue triage.",
    "issue_resolution": "No issues reported during the hypercare window.",
    "handover": "Operational ownership transferred to the Platform Administrator role.",
    "lessons_learned": "Early upstream-context plumbing paid off across every later phase.",
    "closure_statement": (
        "No unresolved critical defects — confirmed via the Test Workbook and carried forward "
        "through the IQ Document. Project is approved for closure."
    ),
}


def _make_request(**overrides) -> AgentRunRequest:
    defaults = dict(
        execution_mode="orchestrated",
        project_id="proj-1",
        invocation_id="inv-1",
        agent_id="hypercare_closure",
        task_id="task-1",
        task_request="Close out the booking portal project",
        lifecycle_phase="hypercare_closure",
        constraints={"project_name": "Facilities Booking Portal"},
        run_number=1,
        upstream_artefacts_text={
            "iq_document": "Verified: the Test Workbook's Defects sheet shows zero open defects."
        },
    )
    defaults.update(overrides)
    return AgentRunRequest(**defaults)


def test_execute_renders_artefact_from_real_model_content():
    adapter = HypercareClosureLlmAdapter(provider=FakeModelProvider(json.dumps(_VALID_RESPONSE)))

    result = adapter.execute(_make_request())

    assert len(result.artefacts_produced) == 1
    produced = result.artefacts_produced[0]
    assert produced.artefact_type == "hypercare_closure_report"
    assert produced.entities == []

    doc = Document(produced.file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Facilities Booking Portal" in full_text
    assert "No unresolved critical defects" in full_text


def test_upstream_iq_document_text_is_included_in_prompt():
    provider = FakeModelProvider(json.dumps(_VALID_RESPONSE))
    adapter = HypercareClosureLlmAdapter(provider=provider)

    adapter.execute(_make_request())

    call = provider.calls[0]
    assert "Hypercare" in call["system"]
    assert "zero open defects" in call["user"]


def test_no_upstream_text_omits_context_gracefully():
    provider = FakeModelProvider(json.dumps(_VALID_RESPONSE))
    adapter = HypercareClosureLlmAdapter(provider=provider)

    adapter.execute(_make_request(upstream_artefacts_text={}))

    assert "Approved" not in provider.calls[0]["user"]


def test_missing_closure_statement_raises_model_provider_error():
    response = dict(_VALID_RESPONSE)
    response["closure_statement"] = ""
    adapter = HypercareClosureLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())


def test_missing_free_text_field_falls_back_gracefully():
    response = dict(_VALID_RESPONSE)
    del response["lessons_learned"]
    adapter = HypercareClosureLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    result = adapter.execute(_make_request())

    doc = Document(result.artefacts_produced[0].file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Lessons learned not specified by the model." in full_text


def test_malformed_json_raises_model_provider_error():
    adapter = HypercareClosureLlmAdapter(provider=FakeModelProvider("not json"))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())
