import json

import pytest
from docx import Document

from app.adapters.llm_agent_adapter import TechnicalDesignLlmAdapter
from app.adapters.model_providers.base import ModelProviderError
from app.agents_registry.contract import AgentRunRequest
from tests.fixtures.fake_model_provider import FakeModelProvider

_VALID_RESPONSE = {
    "options": [
        {"name": "Model-driven app on Dataverse", "tradeoff": "Faster to build, less UI flexibility."},
        {"name": "Canvas app on Dataverse", "tradeoff": "Full UI control, more build effort."},
    ],
    "architecture_decisions": [
        "Use Dataverse as the system of record for all request data.",
        "Use Power Automate for all cross-entity workflow orchestration.",
        "Expose external system access only through custom connectors.",
    ],
    "risks": ["Dataverse API limits may throttle bulk operations during peak submission periods."],
    "limitations": ["Offline mobile support is out of scope for the initial release."],
    "dependencies": ["Requires a Dataverse environment provisioned before the Build phase starts."],
    "logical_architecture": "Power Platform canvas app frontend, Dataverse system of record.",
    "integration_overview": "External systems are integrated via dedicated custom connectors.",
    "infrastructure_overview": "Dev/test/prod Power Platform environments with solution-based ALM.",
}


def _make_request(**overrides) -> AgentRunRequest:
    defaults = dict(
        execution_mode="orchestrated",
        project_id="proj-1",
        invocation_id="inv-1",
        agent_id="technical_design",
        task_id="task-1",
        task_request="Design the technical architecture for onboarding",
        lifecycle_phase="technical_design",
        constraints={"project_name": "Employee Onboarding"},
        run_number=1,
        upstream_artefacts_text={"ux_design_specification": "SCR-001: Dashboard — summarizes open items."},
    )
    defaults.update(overrides)
    return AgentRunRequest(**defaults)


def test_execute_renders_both_artefacts_from_real_model_content():
    adapter = TechnicalDesignLlmAdapter(provider=FakeModelProvider(json.dumps(_VALID_RESPONSE)))

    result = adapter.execute(_make_request())

    assert len(result.artefacts_produced) == 2
    types = {p.artefact_type for p in result.artefacts_produced}
    assert types == {"solution_approach", "architecture_handbook"}

    solution_approach = next(p for p in result.artefacts_produced if p.artefact_type == "solution_approach")
    assert solution_approach.entities == ["ADR-001", "ADR-002", "ADR-003"]
    doc = Document(solution_approach.file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Employee Onboarding" in full_text
    assert "ADR-001: Use Dataverse as the system of record for all request data." in full_text
    assert "Model-driven app on Dataverse" in full_text

    handbook = next(p for p in result.artefacts_produced if p.artefact_type == "architecture_handbook")
    assert handbook.entities == []
    handbook_doc = Document(handbook.file_path)
    handbook_text = "\n".join(p.text for p in handbook_doc.paragraphs)
    assert "Employee Onboarding" in handbook_text
    assert "Dataverse system of record" in handbook_text


def test_upstream_ux_design_text_is_included_in_prompt():
    provider = FakeModelProvider(json.dumps(_VALID_RESPONSE))
    adapter = TechnicalDesignLlmAdapter(provider=provider)

    adapter.execute(_make_request())

    call = provider.calls[0]
    assert "Technical Design Agent" in call["system"]
    assert "SCR-001: Dashboard — summarizes open items." in call["user"]


def test_missing_upstream_text_omits_context_section_gracefully():
    provider = FakeModelProvider(json.dumps(_VALID_RESPONSE))
    adapter = TechnicalDesignLlmAdapter(provider=provider)

    adapter.execute(_make_request(upstream_artefacts_text={}))

    assert "Approved UX Design Specification" not in provider.calls[0]["user"]


def test_strips_model_supplied_decision_id_prefix():
    response = dict(_VALID_RESPONSE)
    response["architecture_decisions"] = ["ADR-1: Use Dataverse as the system of record."]
    adapter = TechnicalDesignLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    result = adapter.execute(_make_request())

    solution_approach = next(p for p in result.artefacts_produced if p.artefact_type == "solution_approach")
    doc = Document(solution_approach.file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "ADR-001: Use Dataverse as the system of record." in full_text
    assert "ADR-001: ADR" not in full_text


@pytest.mark.parametrize(
    "missing_key,missing_value",
    [("options", []), ("architecture_decisions", [])],
)
def test_missing_required_field_raises_model_provider_error(missing_key, missing_value):
    response = dict(_VALID_RESPONSE)
    response[missing_key] = missing_value
    adapter = TechnicalDesignLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())


def test_single_option_raises_model_provider_error():
    response = dict(_VALID_RESPONSE)
    response["options"] = [{"name": "Only option", "tradeoff": "No comparison offered."}]
    adapter = TechnicalDesignLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())


def test_malformed_json_raises_model_provider_error():
    adapter = TechnicalDesignLlmAdapter(provider=FakeModelProvider("not json"))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())
