import json

import pytest
from docx import Document

from app.adapters.llm_agent_adapter import BuildLlmAdapter
from app.adapters.model_providers.base import ModelProviderError
from app.agents_registry.contract import AgentRunRequest
from tests.fixtures.fake_model_provider import FakeModelProvider

_VALID_RESPONSE = {
    "implementation_assets": "Canvas app screens, Dataverse solution, and Power Automate flows built.",
    "configuration_summary": "Environment variables and connection references configured per governance.",
    "findings": [
        {"reference": "SCR-002", "description": "Request Form screen was missing client-side validation."},
        {"reference": "ADR-001", "description": "Dataverse concurrency control was not yet enabled."},
    ],
}


def _make_request(**overrides) -> AgentRunRequest:
    defaults = dict(
        execution_mode="orchestrated",
        project_id="proj-1",
        invocation_id="inv-1",
        agent_id="build",
        task_id="task-1",
        task_request="Build the booking portal",
        lifecycle_phase="build",
        constraints={"project_name": "Facilities Booking Portal"},
        run_number=1,
        upstream_artefacts_text={
            "ux_design_specification": "SCR-002: Request Form — guided form for submitting a request.",
            "solution_approach": "ADR-001: Use Dataverse as the system of record.",
            "governance_document": "DLP: Dataverse business data group.",
        },
    )
    defaults.update(overrides)
    return AgentRunRequest(**defaults)


def test_execute_renders_both_artefacts_from_real_model_content():
    adapter = BuildLlmAdapter(provider=FakeModelProvider(json.dumps(_VALID_RESPONSE)))

    result = adapter.execute(_make_request())

    assert len(result.artefacts_produced) == 2
    types = {p.artefact_type for p in result.artefacts_produced}
    assert types == {"build_review_report", "final_code_review_report"}

    build_review = next(p for p in result.artefacts_produced if p.artefact_type == "build_review_report")
    assert build_review.entities == ["DEF-001", "DEF-002"]
    doc = Document(build_review.file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Facilities Booking Portal" in full_text
    assert "DEF-001: SCR-002: Request Form screen was missing client-side validation." in full_text

    final_review = next(p for p in result.artefacts_produced if p.artefact_type == "final_code_review_report")
    assert final_review.entities == ["DEF-001", "DEF-002"]


def test_multiple_upstream_artefacts_are_included_in_prompt():
    provider = FakeModelProvider(json.dumps(_VALID_RESPONSE))
    adapter = BuildLlmAdapter(provider=provider)

    adapter.execute(_make_request())

    call = provider.calls[0]
    assert "Build Agent" in call["system"]
    assert "SCR-002: Request Form" in call["user"]
    assert "ADR-001: Use Dataverse" in call["user"]
    assert "DLP: Dataverse business data group." in call["user"]


def test_no_upstream_text_omits_context_gracefully():
    provider = FakeModelProvider(json.dumps(_VALID_RESPONSE))
    adapter = BuildLlmAdapter(provider=provider)

    adapter.execute(_make_request(upstream_artefacts_text={}))

    call = provider.calls[0]
    assert "Approved" not in call["user"]


def test_missing_findings_raises_model_provider_error():
    response = dict(_VALID_RESPONSE)
    response["findings"] = []
    adapter = BuildLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())


def test_finding_without_reference_falls_back_to_description_only():
    response = dict(_VALID_RESPONSE)
    response["findings"] = [{"reference": "", "description": "General code quality issue."}]
    adapter = BuildLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    result = adapter.execute(_make_request())

    doc = Document(result.artefacts_produced[0].file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "DEF-001: General code quality issue." in full_text


def test_malformed_json_raises_model_provider_error():
    adapter = BuildLlmAdapter(provider=FakeModelProvider("not json"))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())
