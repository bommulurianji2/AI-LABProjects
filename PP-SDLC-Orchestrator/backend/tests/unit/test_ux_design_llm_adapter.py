import json

import pytest
from docx import Document

from app.adapters.llm_agent_adapter import UxDesignLlmAdapter
from app.adapters.model_providers.base import ModelProviderError
from app.agents_registry.contract import AgentRunRequest
from tests.fixtures.fake_model_provider import FakeModelProvider

_VALID_RESPONSE = {
    "personas": [
        "Priya, an HR administrator who processes requests in bulk.",
        "Sam, a first-time employee user.",
    ],
    "journeys": [
        "Submit a request and track its status to resolution.",
        "Review a pending request and approve or reject it.",
    ],
    "screens": [
        {"name": "Dashboard", "description": "Summarizes open items."},
        {"name": "Request Form", "description": "Guided form for submitting a request."},
        {"name": "Approvals Queue", "description": "List of items awaiting decision."},
    ],
    "navigation": "Top-level tabs for Dashboard, Requests, and Approvals.",
    "responsive_behavior": "Collapses to a single column below 768px.",
    "accessibility": "All controls are keyboard-reachable with WCAG AA contrast.",
}


def _make_request(**overrides) -> AgentRunRequest:
    defaults = dict(
        execution_mode="orchestrated",
        project_id="proj-1",
        invocation_id="inv-1",
        agent_id="ux_design",
        task_id="task-1",
        task_request="Design the onboarding experience",
        lifecycle_phase="ux_design",
        constraints={"project_name": "Employee Onboarding"},
        run_number=1,
        upstream_artefacts_text={"requirement_specification": "REQ-001: The system shall do a thing."},
    )
    defaults.update(overrides)
    return AgentRunRequest(**defaults)


def test_execute_renders_both_artefacts_from_real_model_content():
    adapter = UxDesignLlmAdapter(provider=FakeModelProvider(json.dumps(_VALID_RESPONSE)))

    result = adapter.execute(_make_request())

    assert len(result.artefacts_produced) == 2
    types = {p.artefact_type for p in result.artefacts_produced}
    assert types == {"ux_design_specification", "ux_interactive_prototype"}
    for produced in result.artefacts_produced:
        assert produced.entities == ["SCR-001", "SCR-002", "SCR-003"]

    spec = next(p for p in result.artefacts_produced if p.artefact_type == "ux_design_specification")
    doc = Document(spec.file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Employee Onboarding" in full_text
    assert "SCR-001: Dashboard" in full_text
    assert "Priya, an HR administrator" in full_text

    prototype = next(p for p in result.artefacts_produced if p.artefact_type == "ux_interactive_prototype")
    html_text = open(prototype.file_path, encoding="utf-8").read()
    assert "Employee Onboarding" in html_text
    assert "SCR-001" in html_text


def test_upstream_requirement_spec_text_is_included_in_prompt():
    provider = FakeModelProvider(json.dumps(_VALID_RESPONSE))
    adapter = UxDesignLlmAdapter(provider=provider)

    adapter.execute(_make_request())

    call = provider.calls[0]
    assert "UX Design Agent" in call["system"]
    assert "REQ-001: The system shall do a thing." in call["user"]


def test_missing_upstream_text_omits_context_section_gracefully():
    provider = FakeModelProvider(json.dumps(_VALID_RESPONSE))
    adapter = UxDesignLlmAdapter(provider=provider)

    adapter.execute(_make_request(upstream_artefacts_text={}))

    assert "Approved Requirement Specification" not in provider.calls[0]["user"]


def test_strips_model_supplied_screen_id_prefix():
    response = dict(_VALID_RESPONSE)
    response["screens"] = [{"name": "SCR-1: Dashboard", "description": "Summary screen."}]
    adapter = UxDesignLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    result = adapter.execute(_make_request())

    spec = next(p for p in result.artefacts_produced if p.artefact_type == "ux_design_specification")
    doc = Document(spec.file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "SCR-001: Dashboard" in full_text
    assert "SCR-001: SCR" not in full_text


@pytest.mark.parametrize(
    "missing_key",
    ["personas", "journeys", "screens"],
)
def test_missing_required_field_raises_model_provider_error(missing_key):
    response = dict(_VALID_RESPONSE)
    response[missing_key] = []
    adapter = UxDesignLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())


def test_malformed_json_raises_model_provider_error():
    adapter = UxDesignLlmAdapter(provider=FakeModelProvider("not json"))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())
