import json

import pytest
from docx import Document

from app.adapters.llm_agent_adapter import GovernanceSecurityLlmAdapter
from app.adapters.model_providers.base import ModelProviderError
from app.agents_registry.contract import AgentRunRequest
from tests.fixtures.fake_model_provider import FakeModelProvider

_VALID_RESPONSE = {
    "identity_design": "Microsoft Entra ID as the identity provider; delegated permissions by default.",
    "permissions": "Least privilege: delegated Graph permissions unless justified otherwise.",
    "environment_strategy": "Separate dev, test, and production Power Platform environments.",
    "dlp": [
        "Dataverse and Office 365 Outlook connector: business data group.",
        "All other connectors: non-business data group, blocked by default.",
    ],
    "connector_governance": "The Office 365 Outlook connector requires explicit DLP approval before use.",
    "licensing": ["Microsoft 365 E3 licenses required for all volunteer and administrator accounts."],
    "compliance": "No regulated data categories identified for this project.",
    "operational_ownership": "Platform Administrator owns environment health; Project Owner owns escalation.",
    "audit_requirements": "All approval and rework events must be captured in the audit log.",
}


def _make_request(**overrides) -> AgentRunRequest:
    defaults = dict(
        execution_mode="orchestrated",
        project_id="proj-1",
        invocation_id="inv-1",
        agent_id="governance_security",
        task_id="task-1",
        task_request="Produce governance for the booking portal",
        lifecycle_phase="governance_security",
        constraints={"project_name": "Facilities Booking Portal"},
        run_number=1,
        upstream_artefacts_text={
            "data_design_document": "DATA-001: Facility. Connector: Office 365 Outlook connector."
        },
    )
    defaults.update(overrides)
    return AgentRunRequest(**defaults)


def test_execute_renders_artefact_from_real_model_content():
    adapter = GovernanceSecurityLlmAdapter(provider=FakeModelProvider(json.dumps(_VALID_RESPONSE)))

    result = adapter.execute(_make_request())

    assert len(result.artefacts_produced) == 1
    produced = result.artefacts_produced[0]
    assert produced.artefact_type == "governance_document"
    assert produced.entities == []

    doc = Document(produced.file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Facilities Booking Portal" in full_text
    assert "Office 365 Outlook connector: business data group." in full_text
    assert "explicit DLP approval before use" in full_text


def test_upstream_data_design_text_is_included_in_prompt():
    provider = FakeModelProvider(json.dumps(_VALID_RESPONSE))
    adapter = GovernanceSecurityLlmAdapter(provider=provider)

    adapter.execute(_make_request())

    call = provider.calls[0]
    assert "Governance & Security Agent" in call["system"]
    assert "DATA-001: Facility." in call["user"]


def test_missing_upstream_text_omits_context_section_gracefully():
    provider = FakeModelProvider(json.dumps(_VALID_RESPONSE))
    adapter = GovernanceSecurityLlmAdapter(provider=provider)

    adapter.execute(_make_request(upstream_artefacts_text={}))

    assert "Approved Data Design Document" not in provider.calls[0]["user"]


@pytest.mark.parametrize("missing_key", ["dlp", "licensing"])
def test_missing_required_list_field_raises_model_provider_error(missing_key):
    response = dict(_VALID_RESPONSE)
    response[missing_key] = []
    adapter = GovernanceSecurityLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())


def test_missing_free_text_field_falls_back_gracefully():
    response = dict(_VALID_RESPONSE)
    del response["compliance"]
    adapter = GovernanceSecurityLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    result = adapter.execute(_make_request())

    doc = Document(result.artefacts_produced[0].file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Compliance not specified by the model." in full_text


def test_malformed_json_raises_model_provider_error():
    adapter = GovernanceSecurityLlmAdapter(provider=FakeModelProvider("not json"))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())
