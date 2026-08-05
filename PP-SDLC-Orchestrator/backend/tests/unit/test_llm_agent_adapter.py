import json

import pytest
from docx import Document

from app.adapters.llm_agent_adapter import AnalysisLlmAdapter
from app.adapters.model_providers.base import ModelProviderError
from app.agents_registry.contract import AgentRunRequest
from tests.fixtures.fake_model_provider import FakeModelProvider


def _make_request(**overrides) -> AgentRunRequest:
    defaults = dict(
        execution_mode="orchestrated",
        project_id="proj-1",
        invocation_id="inv-1",
        agent_id="analysis",
        task_id="task-1",
        task_request="Build an employee leave request workflow with manager approval",
        lifecycle_phase="analysis",
        constraints={"project_name": "Employee Leave Request"},
        run_number=1,
    )
    defaults.update(overrides)
    return AgentRunRequest(**defaults)


def test_execute_renders_real_model_content_into_docx():
    provider = FakeModelProvider(
        json.dumps(
            {
                "scope": "Submitting and approving employee leave requests.",
                "functional_requirements": [
                    "The system shall let an employee submit a leave request with start and end dates.",
                    "The system shall route a submitted request to the employee's manager for approval.",
                ],
                "assumptions": ["Managers are already known to the system via Entra ID."],
            }
        )
    )
    adapter = AnalysisLlmAdapter(provider=provider)

    result = adapter.execute(_make_request())

    assert len(result.artefacts_produced) == 1
    produced = result.artefacts_produced[0]
    assert produced.entities == ["REQ-001", "REQ-002"]

    doc = Document(produced.file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Employee Leave Request" in full_text
    assert "REQ-001: The system shall let an employee submit" in full_text
    assert "REQ-002: The system shall route a submitted request" in full_text
    assert "Submitting and approving employee leave requests." in full_text
    assert "Managers are already known to the system via Entra ID." in full_text


def test_execute_passes_skill_md_as_system_prompt_and_task_request_in_user_prompt():
    provider = FakeModelProvider()
    adapter = AnalysisLlmAdapter(provider=provider)

    adapter.execute(_make_request(task_request="A very specific task request marker"))

    assert len(provider.calls) == 1
    call = provider.calls[0]
    assert "Analysis Agent" in call["system"]  # from the real SKILL.md
    assert "A very specific task request marker" in call["user"]
    assert "JSON" in call["user"]


def test_handles_json_wrapped_in_markdown_code_fence():
    fenced = "Here you go:\n```json\n" + json.dumps(
        {
            "scope": "Scope text.",
            "functional_requirements": ["The system shall do a thing."],
            "assumptions": [],
        }
    ) + "\n```"
    adapter = AnalysisLlmAdapter(provider=FakeModelProvider(fenced))

    result = adapter.execute(_make_request())

    assert result.artefacts_produced[0].entities == ["REQ-001"]


def test_malformed_json_raises_model_provider_error():
    adapter = AnalysisLlmAdapter(provider=FakeModelProvider("this is not json at all"))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())


def test_empty_functional_requirements_raises_model_provider_error():
    empty = json.dumps({"scope": "Scope.", "functional_requirements": [], "assumptions": []})
    adapter = AnalysisLlmAdapter(provider=FakeModelProvider(empty))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())


@pytest.mark.parametrize(
    "model_supplied_text",
    [
        "REQ-001: The system shall let a user do a thing.",
        "REQ-1: The system shall let a user do a thing.",
        "REQ001: The system shall let a user do a thing.",
        "1. The system shall let a user do a thing.",
        "1) The system shall let a user do a thing.",
    ],
)
def test_strips_model_supplied_id_prefix_to_avoid_doubling_up(model_supplied_text):
    """Regression test: found via manual testing that some models prepend
    their own numbering (e.g. "REQ-001: ...") despite being told not to,
    which then got double-prefixed by the system's own REQ-00N assignment,
    producing "REQ-001: REQ-001: The system shall...".
    """
    response = json.dumps(
        {"scope": "Scope.", "functional_requirements": [model_supplied_text], "assumptions": []}
    )
    adapter = AnalysisLlmAdapter(provider=FakeModelProvider(response))

    result = adapter.execute(_make_request())

    doc = Document(result.artefacts_produced[0].file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "REQ-001: The system shall let a user do a thing." in full_text
    assert "REQ-001: REQ" not in full_text


def test_missing_assumptions_key_defaults_gracefully():
    no_assumptions = json.dumps(
        {"scope": "Scope.", "functional_requirements": ["The system shall do a thing."]}
    )
    adapter = AnalysisLlmAdapter(provider=FakeModelProvider(no_assumptions))

    result = adapter.execute(_make_request())

    doc = Document(result.artefacts_produced[0].file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "No assumptions provided by the model." in full_text
