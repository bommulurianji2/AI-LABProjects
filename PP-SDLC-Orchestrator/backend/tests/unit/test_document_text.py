from pathlib import Path

from docx import Document
from openpyxl import Workbook

from app.adapters.document_text import extract_text


def test_extract_text_reads_docx_paragraphs(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("Hello world")
    doc.add_paragraph("")
    doc.add_paragraph("Second line")
    file_path = tmp_path / "doc.docx"
    doc.save(file_path)

    text = extract_text(file_path)

    assert "Hello world" in text
    assert "Second line" in text


def test_extract_text_reads_xlsx_rows_across_sheets(tmp_path: Path):
    wb = Workbook()
    ws = wb.active
    ws.title = "Test Cases"
    ws.append(["Test ID", "Type", "Status"])
    ws.append(["TC-001", "OQ", "Passed"])
    defects_ws = wb.create_sheet("Defects")
    defects_ws.append(["Defect ID", "Status"])
    defects_ws.append(["DEF-001", "Open"])
    file_path = tmp_path / "workbook.xlsx"
    wb.save(file_path)

    text = extract_text(file_path)

    assert "[Test Cases]" in text
    assert "TC-001 | OQ | Passed" in text
    assert "[Defects]" in text
    assert "DEF-001 | Open" in text


def test_extract_text_truncates_to_max_chars(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("x" * 100)
    file_path = tmp_path / "doc.docx"
    doc.save(file_path)

    text = extract_text(file_path, max_chars=10)

    assert len(text) == 10


def test_extract_text_falls_back_to_raw_read_for_unknown_extensions(tmp_path: Path):
    file_path = tmp_path / "notes.txt"
    file_path.write_text("plain text content", encoding="utf-8")

    text = extract_text(file_path)

    assert text == "plain text content"
