import json

import pytest
from docx import Document

from app.adapters.llm_agent_adapter import ValidationQaLlmAdapter
from app.adapters.model_providers.base import ModelProviderError
from app.agents_registry.contract import AgentRunRequest
from tests.fixtures.fake_model_provider import FakeModelProvider

_VALID_RESPONSE = {
    "validation_scope": "Independent validation of all approved design and build artefacts.",
    "standards_assessment": "Assessed against accessibility, naming convention, and traceability standards.",
    "findings": [
        {
            "reference": "DEF-002",
            "description": "The delegation warning on My Reservations was not re-tested after the fix.",
            "remediation_owner": "Build Agent",
        },
        {
            "reference": "ADR-001",
            "description": "Dataverse concurrency control lacks a documented rollback plan.",
            "remediation_owner": "Technical Design Agent",
        },
    ],
    "overall_verdict": "Pass with findings — no critical defects block progression.",
}


def _make_request(**overrides) -> AgentRunRequest:
    defaults = dict(
        execution_mode="orchestrated",
        project_id="proj-1",
        invocation_id="inv-1",
        agent_id="validation_qa",
        task_id="task-1",
        task_request="Validate the booking portal build",
        lifecycle_phase="validation_qa",
        constraints={"project_name": "Facilities Booking Portal"},
        run_number=1,
        upstream_artefacts_text={
            "final_code_review_report": "DEF-002: resolved.",
            "solution_approach": "ADR-001: Use Dataverse as the system of record.",
        },
    )
    defaults.update(overrides)
    return AgentRunRequest(**defaults)


def test_execute_renders_artefact_from_real_model_content():
    adapter = ValidationQaLlmAdapter(provider=FakeModelProvider(json.dumps(_VALID_RESPONSE)))

    result = adapter.execute(_make_request())

    assert len(result.artefacts_produced) == 1
    produced = result.artefacts_produced[0]
    assert produced.artefact_type == "validation_report"
    assert produced.entities == []

    doc = Document(produced.file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Facilities Booking Portal" in full_text
    assert "DEF-002: The delegation warning on My Reservations was not re-tested after the fix. (Remediation owner: Build Agent)" in full_text
    assert "Pass with findings" in full_text


def test_multiple_upstream_artefacts_are_included_in_prompt():
    provider = FakeModelProvider(json.dumps(_VALID_RESPONSE))
    adapter = ValidationQaLlmAdapter(provider=provider)

    adapter.execute(_make_request())

    call = provider.calls[0]
    assert "Validation" in call["system"]
    assert "DEF-002: resolved." in call["user"]
    assert "ADR-001: Use Dataverse" in call["user"]


def test_no_upstream_text_omits_context_gracefully():
    provider = FakeModelProvider(json.dumps(_VALID_RESPONSE))
    adapter = ValidationQaLlmAdapter(provider=provider)

    adapter.execute(_make_request(upstream_artefacts_text={}))

    assert "Approved" not in provider.calls[0]["user"]


def test_missing_findings_raises_model_provider_error():
    response = dict(_VALID_RESPONSE)
    response["findings"] = []
    adapter = ValidationQaLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())


def test_finding_without_remediation_owner_is_dropped():
    response = dict(_VALID_RESPONSE)
    response["findings"] = [
        {"reference": "DEF-002", "description": "Missing owner.", "remediation_owner": ""},
        {"reference": "ADR-001", "description": "Has an owner.", "remediation_owner": "Technical Design Agent"},
    ]
    adapter = ValidationQaLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    result = adapter.execute(_make_request())

    doc = Document(result.artefacts_produced[0].file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Missing owner." not in full_text
    assert "Has an owner." in full_text


def test_all_findings_without_remediation_owner_raises():
    response = dict(_VALID_RESPONSE)
    response["findings"] = [{"reference": "DEF-002", "description": "No owner given.", "remediation_owner": ""}]
    adapter = ValidationQaLlmAdapter(provider=FakeModelProvider(json.dumps(response)))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())


def test_malformed_json_raises_model_provider_error():
    adapter = ValidationQaLlmAdapter(provider=FakeModelProvider("not json"))

    with pytest.raises(ModelProviderError):
        adapter.execute(_make_request())
