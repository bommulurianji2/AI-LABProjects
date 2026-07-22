"""Extends the chain to a sixth real agent: ... -> Governance & Security ->
Build. Confirms both build artefacts are generated with matching seeded
DEF-00N entities and approval unlocks validation_qa.
"""

from docx import Document

from app.domain.enums import ArtefactVersionStatus, PhaseStatus
from app.models.artefact import ArtefactVersion
from app.models.user import User
from tests.helpers import make_orchestrator, run_phase_to_approval


def test_chain_through_build_unlocks_validation_qa(db_session):
    orchestrator = make_orchestrator(db_session)
    reviewer = User(email="developer@example.test", role="Contributor")
    db_session.add(reviewer)
    db_session.commit()

    project = orchestrator.create_project("Contract and Approval Management")

    run_phase_to_approval(orchestrator, project, reviewer_id=reviewer.id, task_request="Draft requirements")
    run_phase_to_approval(orchestrator, project, reviewer_id=reviewer.id, task_request="Design the experience")
    run_phase_to_approval(
        orchestrator, project, reviewer_id=reviewer.id, task_request="Design the technical architecture"
    )
    run_phase_to_approval(orchestrator, project, reviewer_id=reviewer.id, task_request="Design the data model")
    run_phase_to_approval(
        orchestrator, project, reviewer_id=reviewer.id, task_request="Define governance and security controls"
    )
    assert project.current_phase == "build"

    build_run = run_phase_to_approval(orchestrator, project, reviewer_id=reviewer.id, task_request="Build the solution")
    assert project.current_phase == "validation_qa"
    assert project.phase_status == PhaseStatus.PENDING.value

    versions = db_session.query(ArtefactVersion).filter_by(run_id=build_run.id).all()
    assert len(versions) == 2
    for v in versions:
        db_session.refresh(v)
        assert v.version_label == "v1.0"
        assert v.status == ArtefactVersionStatus.BASELINE.value

    build_review_version = next(v for v in versions if "build_review_report" in v.file_path)
    doc = Document(build_review_version.file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "DEF-001" in full_text

    final_review_version = next(v for v in versions if "final_code_review_report" in v.file_path)
    final_doc = Document(final_review_version.file_path)
    final_text = "\n".join(p.text for p in final_doc.paragraphs)
    assert "DEF-001: resolved." in final_text
