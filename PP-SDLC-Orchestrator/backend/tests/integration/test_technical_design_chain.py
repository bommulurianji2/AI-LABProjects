"""Extends the chain to a third real agent: Analysis -> UX Design ->
Technical Design. Confirms the Technical Design Agent's two artefacts
(Solution Approach + Architecture Handbook) are generated with correct
seeded ADR entities and that approval unlocks data_integration.
"""

from docx import Document

from app.domain.enums import ArtefactVersionStatus, PhaseStatus, ReviewDecision
from app.models.artefact import ArtefactVersion
from app.models.user import User
from tests.helpers import make_orchestrator, run_phase_to_approval


def test_chain_through_technical_design_unlocks_data_integration(db_session):
    orchestrator = make_orchestrator(db_session)
    reviewer = User(email="solution.architect@example.test", role="Reviewer")
    db_session.add(reviewer)
    db_session.commit()

    project = orchestrator.create_project("Contract and Approval Management")

    run_phase_to_approval(orchestrator, project, reviewer_id=reviewer.id, task_request="Draft requirements")
    assert project.current_phase == "ux_design"

    run_phase_to_approval(orchestrator, project, reviewer_id=reviewer.id, task_request="Design the experience")
    assert project.current_phase == "technical_design"

    td_run = run_phase_to_approval(
        orchestrator, project, reviewer_id=reviewer.id, task_request="Design the technical architecture"
    )
    assert project.current_phase == "data_integration"
    assert project.phase_status == PhaseStatus.PENDING.value

    versions = db_session.query(ArtefactVersion).filter_by(run_id=td_run.id).all()
    assert len(versions) == 2
    for v in versions:
        db_session.refresh(v)
        assert v.version_label == "v1.0"
        assert v.status == ArtefactVersionStatus.BASELINE.value

    solution_approach_version = next(v for v in versions if "solution_approach" in v.file_path)
    doc = Document(solution_approach_version.file_path)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "ADR-001" in full_text
    assert "Contract and Approval Management" in full_text
    assert "Recommended:" in full_text

    handbook_version = next(v for v in versions if "architecture_handbook" in v.file_path)
    handbook_doc = Document(handbook_version.file_path)
    handbook_text = "\n".join(p.text for p in handbook_doc.paragraphs)
    assert "Dataverse" in handbook_text


def test_rejected_technical_design_keeps_project_on_phase(db_session):
    orchestrator = make_orchestrator(db_session)
    reviewer = User(email="solution.architect@example.test", role="Reviewer")
    db_session.add(reviewer)
    db_session.commit()

    project = orchestrator.create_project("Regulated Document Review")
    run_phase_to_approval(orchestrator, project, reviewer_id=reviewer.id, task_request="Draft requirements")
    run_phase_to_approval(orchestrator, project, reviewer_id=reviewer.id, task_request="Design the experience")

    run = orchestrator.start_run(project, task_request="Design the technical architecture")
    orchestrator.submit_review(run, reviewer_id=reviewer.id, decision=ReviewDecision.REJECTED)

    assert project.current_phase == "technical_design"
    assert project.phase_status == PhaseStatus.PENDING.value
